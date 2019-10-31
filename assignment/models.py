from django.db import models
from course.models import Course
from django.urls import reverse
from authentication.models import User
from docx.enum.style import WD_STYLE_TYPE
from django.conf import settings

import os
import sys
import subprocess
import docx
from docx.shared import Pt
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Image,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor
from io import BytesIO


# Create your models here.

def get_upload_location(instance, filename):
    assignment_name = f"assignment_{instance.question.assignment.id}"
    course_name = f"course_{instance.question.assignment.course.id}"
    return os.path.join(
        'assignments',
        os.path.join(
            course_name,
            os.path.join(
                assignment_name,
                filename
            )
        )
    )


def add_header(input_file, header_text):
    input_doc = docx.Document(input_file)
    header = input_doc.sections[0].header
    for run in header.paragraphs[0].runs:
        run.text = ''
    run = header.paragraphs[0].add_run(f'Question: {header_text}')
    run.font.color.rgb = docx.shared.RGBColor(0x3b, 0x3b, 0x3b)
    run.font.size = Pt(13)
    header.paragraphs[0].style = input_doc.styles['Heading']
    if len(header.paragraphs) == 1:
        header.add_paragraph()
    input_doc.save(input_file)


def convert_image_word(input_file, output_file_dir, file_name):
    doc = docx.Document()
    doc.add_picture(input_file, width=docx.shared.Inches(6))
    doc.styles.add_style("Heading", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    doc_path = os.path.join(output_file_dir, f'{file_name}.docx')
    doc.save(doc_path)

    return doc_path


def convert_image_pdf(input_file, output_file, header_text):

    pdf = SimpleDocTemplate(output_file)

    sample_style_sheet = getSampleStyleSheet()

    parastyle = ParagraphStyle(
        'header',
        fontSize=13,
        textColor=HexColor(0x3b3b3b)
        # fontName="Roboto"
    )

    flowables = [
        Paragraph(f'Question: {header_text}', parastyle,),
        Spacer(2, height=2*cm),
        Image(
            input_file,
            width=6*inch,
            height=6*inch,
            kind='proportional'
        ),
    ]

    pdf.build(flowables=flowables)


def convert_word_pdf(input_file, output_file, header_text):

    add_header(input_file, header_text)

    args = [
        'libreoffice',
        '--headless',
        '--convert-to',
        'pdf',
        '--outdir',
        output_file, input_file,
    ]
    process = subprocess.run(
        args,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=10
    )


class Assignment(models.Model):
    course = models.ForeignKey(
        Course, on_delete=models.CASCADE, related_name='assignments')
    name = models.CharField(max_length=20, null=False, blank=False)
    description = models.CharField(
        max_length=500,
        null=False,
        blank=False,
        default='no description provided by tutor',
    )
    deadline = models.DateField(null=False, blank=False)
    date_created = models.DateTimeField(null=False, auto_now_add=True)

    def get_absolute_url(self):
        return reverse('assignment:create-question', kwargs={
            "course_id": self.course.id,
            "assignment_id": self.id,
            "course_type": "my-courses"
        })

    def __str__(self):
        return f'course {self.course.id} assignment: {self.id}'


class Question(models.Model):
    question = models.CharField(max_length=1000, blank=False, null=False)
    assignment = models.ForeignKey(
        Assignment, on_delete=models.CASCADE, related_name='questions')

    def __str__(self):
        return f'assignment {self.assignment.id} question {self.id}'


class Submission(models.Model):
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='all_submissions'
    )
    question = models.ForeignKey(
        Question,
        on_delete=models.CASCADE,
        related_name='submissions'
    )
    solution = models.FileField(
        upload_to=get_upload_location,
        null=True
    )
    solution_pdf = models.FileField(
        null=True
    )
    date_created = models.DateTimeField(auto_now_add=True, null=False)

    def add_pdf_solution(self):
        input_file = self.solution.path
        file_name = os.path.basename(input_file)
        file_format = file_name.rsplit('.')[-1]
        output_file_dir = os.path.dirname(input_file)
        output_file = input_file.replace(f'.{file_format}', '.pdf')
        input_url = self.solution.url
        output_url = input_url.replace(f'.{file_format}', '.pdf')
        output_url = output_url.replace(settings.MEDIA_URL, '')

        timeout = 10

        image_file_format = ['jpg', 'jpeg', 'png']
        document_file_format = ['doc', 'docx']
        supported_file_format = image_file_format + document_file_format

        is_image = True if file_format in image_file_format else False

        if file_format == 'pdf':
            self.solution_pdf = self.solution
            return
        elif file_format in document_file_format:
            convert_word_pdf(
                input_file,
                output_file_dir,
                self.question.question,
            )
            self.solution_pdf.name = output_url
        elif file_format in image_file_format:
            convert_image_pdf(
                input_file,
                output_file,
                self.question.question
            )
            self.solution_pdf.name = output_url
        else:
            pass

    def save(self, *args, add_pdf=True, **kwargs):
        super().save()
        if add_pdf and self.solution != None:
            self.add_pdf_solution()
            self.save(add_pdf=False)
            return
        return

    def __str__(self):
        return f"solution to question {self.question.id} by {self.user.username}"
